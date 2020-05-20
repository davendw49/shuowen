import re
import sys
import pydocx
import codecs
import pymysql
import pypandoc
import argparse
import os
import base64
import imghdr
import shutil
from wand.image import Image
from tqdm import tqdm
from docx import Document
from docx.shared import Inches, Cm
from bs4 import BeautifulSoup

def convert_to_jpeg(picture):
    '''
    convert non-support image into jpg form
    '''
    img = Image(filename=picture)
    img.format = 'jpeg'
    fn = picture.split(".")[0]+".jpg"
    img.save(filename=fn)
    return fn

def p_discriminate(p):
    '''
    judge whether the line input in a picture 
    '''
    if str(p).find("img") != -1:
        return (0, "pic")
    else:
        return (1, "book")
    
def base64_to_img(b64_data, image_type, order):
    '''
    Convert the base64 data back to regular binary image data 
    and figure out the image type (png, gif, jpg, etc)
    '''
    image_data = base64.b64decode(b64_data)
    # image_type = imghdr.what('', image_data)
    # Create the image file and tell the user about it
    destination_file_name ="img/" + str(order) + "." + str(image_type)
    try:
        destination = open(destination_file_name, 'wb')
    except IOError:
        print ("Unable to create image file. You might not have permission to create a file in this location.")
        exit()
    else:
        destination.write(image_data)
        destination.close()
        return destination_file_name
        # print ("New image file: {}".format(destination_file_name))
        
def re_the_src(img_src):
    '''
    return the src.sessions of the string
    '''
    # print(img_src)
    this_img = BeautifulSoup(img_src, features="lxml").find('img')
    src = str(this_img.get('src'))
    height = str(this_img.get('height'))
    width = str(this_img.get('width'))
    # print(src)
    b64_data = src.split("base64,")[1]
    image_type = src.split("base64,")[0].split("/")[1][:-1]
    return b64_data, image_type, height, width

class tuban():
    '''
    a tuban class represents a tuban's basic structural information including imgs and the number of tuban
    '''
    def __init__(self, input_file):
        '''
        init the class
        para: input_file: the file need to be rearrange 
        pare: output_file: the file after rearrange
        '''
        self.tuban_html = pydocx.PyDocX.to_html(input_file)
        self.tuban_soup = BeautifulSoup(self.tuban_html, features="lxml")
        self.input = input_file
        self.output = "result.html"
        self.errors = []
        self.character = {}
        self.ch_dict = {}
        self.p_list = self.tuban_soup.select('p')
        if os.path.exists("img/"):
            shutil.rmtree("img/", True)
        os.mkdir("img/")
            
    def parse(self):
        '''
        parse the file (html form)
        para: self
        '''
        print("parsing")
        p_list = self.p_list
        flag = 0
        # if the pre_stage = 1 means that the situation is booknum, 0 means the pictures
        pre_stage = 0
        pre_tag = 0
        this_booknum = ""
        this_character = []
        pre_num = ""
        for p in tqdm(p_list):
            if flag%2 == 0:
                this_booknum = ""
            else:
                this_character = []
            if flag%2 == 0:
                stage = p_discriminate(p)[0]
                if stage == 1:
                    if stage != pre_stage:
                        this_booknum = p.text
                        pre_stage = stage
                    else:
                        self.errors.append(("error 1 多重简号或者换行符有误", pre_num, p.text, "flag="+str(flag)))
                        return flag
                else:
                    self.errors.append(("error 2 缺省简号或者字体截图中换行符有误", pre_num, p.text, "flag="+str(flag)))
                    return flag
            else:
                stage = p_discriminate(p)[0]
                if stage == 0:
                    children = p.children
                    tag = 1 # 0: character, 1:image
                    cflag = 0
                    for child in children:
                        # if str(child)
                        if cflag%2 ==0:
                            if str(child).find("img") == -1:
                                tag = 0
                                pre_tag = tag
                                tmp_pair_character = str(child).strip()
                            else:
                                self.errors.append(("error 3 出现连续文字结点", this_booknum, p.text, "flag="+str(flag)))
                                return flag
                        else:
                            if str(child).find("img") != -1:
                                tag = 1
                                if tag != pre_tag:
                                    tmp_pair_image = str(child)
                                else:
                                    self.errors.append(("error 4 出现连续图片", this_booknum, p.text, "flag="+str(flag)))
                                    return flag
                            else:
                                self.errors.append(("error 5 出现连续文字结点", this_booknum, str(child), "flag="+str(flag)))
                                return flag
                
                            if len(tmp_pair_character)!=0 and len(tmp_pair_image)!=0:
                                this_character.append((tmp_pair_character, tmp_pair_image))
                        cflag+=1
                    pre_stage = stage
                else:
                    self.errors.append(("error 6 多重简号或者换行符有误", this_booknum, p.text, "flag="+str(flag)))
                    return flag
                    # flag-=1
            
            if len(this_booknum)!=0 and len(this_character)!=0 and flag%2==1:
                # print(flag, len(this_booknum),len(this_character))
                self.character[this_booknum] = this_character
            pre_num = this_booknum
            flag += 1
    
    def array2dict(self):
        print("array2dict")
        array = self.character
        self.ch_dict = {}
        for k in array.keys():
            for item in array[k]:
                if item[0] not in self.ch_dict.keys():
                    self.ch_dict[item[0]] = []
                    self.ch_dict[item[0]].append((item[1], k))
                else:
                    self.ch_dict[item[0]].append((item[1], k))
    
    def array2dict_save(self):
        print("array2dict_save")
        array = self.character
        self.ch_dict = {}
        flag = 0
        for k in tqdm(array.keys()):
            for item in array[k]:
                if item[0] not in self.ch_dict.keys():
                    self.ch_dict[item[0]] = []
                    res = re_the_src(item[1])
                    img_n = base64_to_img(res[0], res[1], flag)
                    self.ch_dict[item[0]].append((img_n, k))
                else:
                    res = re_the_src(item[1])
                    img_n = base64_to_img(res[0], res[1], flag)
                    self.ch_dict[item[0]].append((img_n, k))
                flag += 1
    
    def get_error(self):
        return self.errors
    
    def get_output_html(self):
        f = open(self.output, "w", encoding="utf8")
        for k in tqdm(self.ch_dict.keys()):
            print("<p>"+k+"</p>", file=f)
            print("<p>", end="", file=f)
            for item in self.ch_dict[k]:
                print(item[0], end="", file=f) # img
                print(item[1], end=",", file=f) # book
            print("</p>", file=f)
        f.close()
        
    def get_output_docx_by_pandoc(self):
        output = pypandoc.convert_file(self.output, 'docx', outputfile="result.docx", extra_args=["-M8GB", "+RTS", "-K4096m", "-RTS"])
    
    def get_output_docx_by_docx(self):
        document = Document()
        for k in tqdm(self.ch_dict.keys()):
            paragraph_character = document.add_paragraph(k)
            paragraph_image = document.add_paragraph()
            for item in self.ch_dict[k]:
                fig = item[0]
                if fig.split(".")[1] != "jpg" or fig.split(".")[1] != "png" or fig.split(".")[1] != "jpeg" :
                    new_fig = convert_to_jpeg(fig)
                    fig = new_fig
                paragraph_image.add_run().add_picture(fig, width=Cm(1.5))
                paragraph_image.add_run(item[1]) # book
        document.save(self.input + '_result.docx')
        
    def clean_cache(self):
        if os.path.exists("img/"):
            shutil.rmtree("img/")
        if os.path.exists("result.html"):
            os.remove("result.html")
        
if __name__ == '__main__':
    
    parser = argparse.ArgumentParser(description='Shuowen Platform.')
    parser.add_argument('-f', '--file', default="../data/test1.docx", help='input wechat address')
    parser.add_argument('-m', '--mode', default="docx", help='return in docx or html')
    args = parser.parse_args()
    
    tb = tuban(args.file)
    try:
        tb.parse()
        tb.array2dict_save()
    except:
        tb.get_error()
    if args.mode == "docx":
        tb.get_output_docx_by_docx()
    elif args.mode == "html":
        tb.get_output_html()
    else:
        print("Wrong Mode")
    
    tb.clean_cache()